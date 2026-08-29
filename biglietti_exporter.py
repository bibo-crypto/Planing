"""Extract order data into the Italian dyeing-ticket workbook and Word form.

The source workbook is the ERP export (``Sheet1``) plus the manually completed
``Doispo-Bagno`` and optional raw-yarn sheet.  The code deliberately matches
headers by normalized text and accepts either one or two header rows.

Reference logic: the machine table, Polmoni rule, Commento split, KG/Densita
formula and Color Tube/VMM22 lookup below were reverse-engineered from the
Power Query ("Densita' Quer", "Magazino", "codes", "PESO ROCCHE") embedded in
a real Delta export -- see the extraction notes shared with the user for the
exact M code these mirror. Anything derived from a file the user has not
uploaded yet (Densita' Query workbook, Magazino Color Tube workbook) degrades
gracefully to a blank cell rather than raising.
"""

from __future__ import annotations

import copy
import re
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

from constants import MACHINE_CAPACITIES


def _clean(v: Any) -> str:
    return " ".join(str(v or "").replace("\xa0", " ").split())


def _key(v: Any) -> str:
    # ``\w`` keeps Arabic headers (e.g. وزن and تحضير خام) as well as
    # the synthetic positional keys used for blank ERP headers.
    return re.sub(r"[^\w]+", " ", _clean(v).lower(), flags=re.UNICODE).strip()


def _number(v: Any) -> float | int | None:
    s = _clean(v).replace(" ", "").replace(",", ".")
    if not s:
        return None
    try:
        n = float(s)
        return int(n) if n.is_integer() else n
    except ValueError:
        return None


def _order_number(v: Any) -> str:
    s = _clean(v)
    if "/" in s:
        s = s.rsplit("/", 1)[-1]
    return s.lstrip("0") or "0"


def _first_match(text: str, pattern: str) -> str:
    m = re.search(pattern, text, re.I)
    return m.group(1) if m else ""


# ---------------------------------------------------------------------------
# M/C (machine) lookup -- Rocche total per Bagno group -> physical machine
# number. Values taken from the reference Power Query's machine table.
# ---------------------------------------------------------------------------
# Mapping remains exporter-specific; the capacity list is centralized.
MACHINE_SIZE_TABLE: dict[int, int] = {
    capacity: machine for capacity, machine in zip(
        MACHINE_CAPACITIES, (11, 12, 9, 10, 7, 8, 5, 6, 3)
    )
}
# Seen only on MED orders in the reference query (7 Rocche -> machine 11).
# Kept separate since it overlaps oddly with the 6 -> 11 rule; flagged for
# the user to confirm once real MED data is available.
MACHINE_SIZE_TABLE_MED_EXTRA: dict[int, int] = {7: 11}

# A nearest-size match is only accepted within this tolerance (fraction of
# the nearest table size); anything further is left blank for a human to
# check, per "لو غير مطابقه بيقربها لأقرب ماكينه ولو بعيد سيبه فاضي".
_MACHINE_MATCH_TOLERANCE = 0.15


def _machine_for_count(count: float | int | None, extra: dict[int, int] | None = None) -> str:
    """Return a machine number; non-positive or missing count means blank."""
    if count is None or count <= 0:
        return ""
    table = dict(MACHINE_SIZE_TABLE)
    if extra:
        table.update(extra)
    if count in table:
        return str(table[count])
    nearest = min(table, key=lambda k: abs(k - count))
    if abs(nearest - count) <= nearest * _MACHINE_MATCH_TOLERANCE:
        return str(table[nearest])
    return ""


def _polmoni_segment(additional: str) -> str:
    """MED only: the 5th dash-separated segment of the additional
    description, kept as-is (e.g. '4 POLMONI') when it mentions Polmoni,
    matching the reference query's own POLMON output column -- this is
    text, not a number, and is what actually gets exported."""
    parts = [p.strip() for p in _clean(additional).split("-")]
    seg = parts[4] if len(parts) > 4 else ""
    return seg if "POLMON" in seg.upper() else ""


def _polmoni_extra(additional: str) -> int:
    """MED only: 'N POLMONI' adds N*4 Rocche worth of machine capacity
    (1 polmone stands in for 4 coni) -- used only for the M/C calculation,
    never exported directly."""
    seg = _polmoni_segment(additional)
    m = re.search(r"(\d+)\s*POLMON", seg, re.I) if seg else None
    return int(m.group(1)) * 4 if m else 0


def _commento_for(customer: str, additional: str) -> str:
    """Per-customer Commento, mirroring the reference query's dash-split
    on ALL '-' occurrences:
    ELVY keeps everything after the 2nd '-' (e.g. 'PG-157474-PO-1466-2026'
    -> 'PO-1466-2026'); MED keeps the 6th segment (.6), which is usually
    absent -> blank, confirmed against real MED data."""
    parts = [p.strip() for p in _clean(additional).split("-")]
    if customer == "ELVY":
        return "-".join(parts[2:]).strip() if len(parts) > 2 else ""
    return parts[5].strip() if len(parts) > 5 else ""


def _segment(additional: str, index: int) -> str:
    """1-based dash-segment of the additional description, 'X' treated as
    unassigned/blank (matches the reference query's own X->null rule)."""
    parts = [p.strip() for p in _clean(additional).split("-")]
    seg = parts[index - 1].strip() if len(parts) >= index else ""
    return "" if seg.upper() == "X" else seg


def _partita_med(additional: str) -> str:
    """Descrizione aggiuntiva ordine.4 -> 'Partita MED' in the reference
    query. Not a PO number for MED orders (that pattern is ELVY-only)."""
    return _segment(additional, 4)


def _cliente_med(additional: str) -> str:
    """Descrizione aggiuntiva ordine.5 -> 'Cliente MED'. The raw segment
    (e.g. 'ORDINE CAM') is expanded through _CLIENTE_MED_EXPANSIONS when
    recognised -- 'ORDINE CAM' -> 'ORDINE CAMPIONARIO' was confirmed
    against 3 real pre-existing tickets in the file the user sent, but
    this is the only code seen so far; ask the user for the rest of the
    mapping if other order-type abbreviations show up blank/unexpanded."""
    seg = _segment(additional, 5)
    return _CLIENTE_MED_EXPANSIONS.get(seg.upper(), seg)


_CLIENTE_MED_EXPANSIONS: dict[str, str] = {
    "ORDINE CAM": "ORDINE CAMPIONARIO",
}


def _po_token(additional: str) -> str:
    """Everything from the word 'PO' onward, e.g. 'PO-1466-2026'.
    ELVY-only naming convention -- MED orders use PM/dispo, not PO."""
    m = re.search(r"\bPO\b[-\s]*[0-9][0-9\-]*", additional, re.I)
    return m.group(0).strip() if m else ""


def _dispo_number(dispo: str) -> str:
    """Numeric part of a Dispo value like 'D-00505450-001' -> '505450'."""
    m = re.search(r"D-0*([0-9]+)", _clean(dispo), re.I)
    if m:
        return m.group(1)
    return _clean(dispo).lstrip("0") or _clean(dispo)


def _dispo_suffix(dispo: str) -> str:
    """The per-row sequence suffix of a Dispo value, e.g.
    'D-00505449-001' -> '1' -- this is what Sheet1's own 'Riga' column
    matches against (confirmed against real data: Riga=1,2,3,4 lines up
    with Dispo-Bagno rows -001,-002,-003,-004 in a different physical
    order than Sheet1 itself). Also accepts a bare number so it can be
    used directly on a Riga value too."""
    s = _clean(dispo)
    m = re.search(r"-0*([0-9]+)\s*$", s)
    if m:
        return str(int(m.group(1)))
    n = _number(s)
    return str(int(n)) if n is not None else ""


def build_output_stem(records: list["OrderRecord"], customer: str) -> str:
    """The identifier used both as the output filename and, for ELVY, as
    the ticket's own Commento -- e.g. 'PO-1466-2026' or 'MED-D-505450-2026'.
    """
    if customer == "ELVY":
        for r in records:
            token = _po_token(r.additional_raw)
            if token:
                return token
        return "ELVY"
    dispo_num = _dispo_number(records[0].dispo) if records else ""
    year = str(datetime.now().year)
    return f"MED-D-{dispo_num}-{year}" if dispo_num else f"MED-{year}"


# ---------------------------------------------------------------------------
# Titolo lookup -- shared with the Situazione tab's "Articoli" upload
# (Articolo Filato -> TITOLO), persisted in situazione_db so uploading the
# file from either tab makes it available to both.
# ---------------------------------------------------------------------------

def load_articoli_titolo_map() -> dict[str, str]:
    try:
        import situazione_db
        return situazione_db.load_codes()  # {articolo_filato: titolo}
    except Exception:
        return {}


def load_articoli_marca_map(path: Path) -> tuple[dict[str, str], list[str]]:
    """Articolo Filato -> Marca (the fuller descriptive text, e.g.
    '100/2 - COTTON 100%') -- confirmed against real data as the actual
    source for the ticket's Titolo, NOT the shorter 'TITOLO' column that
    Situazione's own Articoli upload uses."""
    errors: list[str] = []
    out: dict[str, str] = {}
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    try:
        target = wb["Sheet2"] if "Sheet2" in wb.sheetnames else wb.active
        for row in _read_sheet_rows(target):
            art = _clean(_get(row, "Articolo Filato")).upper()
            marca = _clean(_get(row, "Marca"))
            if art and marca:
                out.setdefault(art, marca)
    finally:
        wb.close()
    if not out:
        errors.append("Nessuna riga Articolo Filato/Marca trovata nel file Articoli.")
    return out, errors


def load_articoli_marca_lookup() -> dict[str, str]:
    """Convenience: re-read whatever Articoli.xlsx was last uploaded via
    the Biglietti tab's own Articoli button (path cached in
    articoli_cache.py), returning {} if none has been uploaded yet."""
    try:
        import articoli_cache
    except Exception:
        return {}
    cache = articoli_cache.load_articoli_cache()
    path = cache.get("source_path")
    if not path or not Path(path).is_file():
        return {}
    marca_map, _errors = load_articoli_marca_map(Path(path))
    return marca_map


def _titolo_lookup(articolo: str, codes_map: dict[str, str]) -> str:
    art = _clean(articolo).upper()
    if not art or not codes_map:
        return ""
    if art in codes_map:
        return codes_map[art]
    if art[:1] in ("C", "G"):
        swapped = ("G" if art[:1] == "C" else "C") + art[1:]
        if swapped in codes_map:
            return codes_map[swapped]
    return ""


# ---------------------------------------------------------------------------
# Prezzo lookup -- reuses the existing Listini/Prezzi infrastructure
# (prezzi_cache remembers the last uploaded file; no separate button needed
# here since the Prezzi tab already provides one).
# ---------------------------------------------------------------------------

def load_prezzo_lookup() -> tuple[dict[tuple, tuple], str]:
    """Returns (lookup, source_file_name). Empty lookup + '' if nothing has
    been uploaded to the Prezzi tab yet."""
    try:
        import prezzi_cache
        import prezzi_logic
    except Exception:
        return {}, ""
    cache = prezzi_cache.load_prezzi_cache()
    path = cache.get("source_path")
    if not path or not Path(path).is_file():
        return {}, ""
    df, _errors = prezzi_logic.load_prezzi(path)
    if df is None or df.empty:
        return {}, ""
    return prezzi_logic.build_price_lookup(df), cache.get("source_file", "")


# ---------------------------------------------------------------------------
# Delivery Date (ELVY only) -- how many working days out an order is quoted,
# based on whether the raw batch/yarn is already assigned (Partita GG) and
# whether it's already priced (Prezzo), plus which machine tier it's on.
# Working days skip Friday (pushed to Saturday if a computed date lands on
# one) -- matches the 'AddDaysSkipFriday' Power Query logic exactly.
# ---------------------------------------------------------------------------
_DELIVERY_MACHINES_LONG = {7, 8, 9, 10, 11, 12}
_DELIVERY_MACHINES_MED = {3, 4, 5, 6}


def _add_days_skip_friday(start_date, days_to_add: int):
    target = start_date + timedelta(days=days_to_add)
    if target.weekday() == 4:  # Friday
        target += timedelta(days=1)
    return target


def _compute_delivery_date(raw_batch: Any, prezzo: Any, machine: Any, today=None) -> Any:
    if not _clean(raw_batch):
        return "Bending for yarn"
    today = today or datetime.now().date()
    try:
        mc = int(_number(machine))
    except (TypeError, ValueError):
        mc = None
    has_prezzo = prezzo not in (None, "")
    if not has_prezzo:
        if mc in _DELIVERY_MACHINES_LONG:
            days = 24
        elif mc in _DELIVERY_MACHINES_MED:
            days = 16
        else:
            days = 8
    else:
        if mc in _DELIVERY_MACHINES_LONG:
            days = 16
        elif mc in _DELIVERY_MACHINES_MED:
            days = 8
        else:
            days = 0
    return _add_days_skip_friday(today, days)


def compute_delivery_date(records: list["OrderRecord"]) -> None:
    today = datetime.now().date()
    for r in records:
        r.delivery_date = _compute_delivery_date(r.raw_batch, r.prezzo, r.machine, today)


def _prezzo_for(articolo: str, codice: str, lookup: dict[tuple, tuple]) -> Any:
    if not lookup:
        return ""
    for key in ((articolo, codice), (_clean(articolo), _clean(codice))):
        if key in lookup:
            return lookup[key][1] or ""
    return ""


# Public alias -- other tabs (Situazione) reuse this same matching logic
# for their own Prezzo column instead of re-implementing it.
prezzo_for = _prezzo_for


# Machines (by their Rocche-based M/C total, e.g. Situazione's own 'mc'
# field or Ordine MED's) that get a $2 surcharge on top of the Listini
# price. Shared here since both Ordine MED's 'PREZZO + 2$' column and
# Situazione's own Prezzo column apply the exact same rule.
PREZZO_SURCHARGE_MACHINES = {24, 32, 56}


def apply_machine_surcharge(price: Any, machine: Any) -> Any:
    """price + 2 when machine is 24/32/56 Rocche, otherwise price
    unchanged. machine may be an int, numeric string, or None/blank --
    anything that doesn't cleanly parse just skips the surcharge rather
    than raising."""
    if not isinstance(price, (int, float)):
        return price
    try:
        mc = int(_number(machine))
    except (TypeError, ValueError):
        return price
    return round(price + 2, 2) if mc in PREZZO_SURCHARGE_MACHINES else price


# ---------------------------------------------------------------------------
# Densita' Query workbook -- KG (PESO ROCCHE query) + Densita`(360-390),
# both keyed by Partita (raw_batch / "Partita GG").
# ---------------------------------------------------------------------------

def load_densita_query(path: Path) -> tuple[dict[int, dict[str, Any]], list[str]]:
    errors: list[str] = []
    out: dict[int, dict[str, Any]] = {}
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    try:
        if "Entry" in wb.sheetnames:
            for row in _read_sheet_rows(wb["Entry"]):
                partita = _number(_get(row, "PARTITA", "Partita"))
                peso_lord = _number(_get(row, "Peso Lord", "Peso Lordo"))
                if partita is None:
                    continue
                peso_net = (peso_lord - 45) / 1000 if peso_lord is not None else None
                out.setdefault(int(partita), {})["peso_net"] = peso_net
                color_tube = _clean(_get(row, "Colore Tube", "Color Tube"))
                if color_tube:
                    out[int(partita)]["color_tube"] = color_tube
        else:
            errors.append('Nel file Densita\' Query.xlsx manca il foglio "Entry" (serve per KG).')
        if "Densita" in wb.sheetnames:
            for row in _read_sheet_rows(wb["Densita"]):
                partita = _number(_get(row, "PARTITA", "Partita"))
                if partita is None:
                    continue
                out.setdefault(int(partita), {})["densita"] = _get(row, "Densita` (360-390)", "Densita")
        else:
            errors.append('Nel file Densita\' Query.xlsx manca il foglio "Densita" (serve per Densita`(360-390)).')
    finally:
        wb.close()
    if not out and not errors:
        errors.append("Nessuna riga valida trovata nel file Densita' Query.")
    return out, errors


# ---------------------------------------------------------------------------
# Magazino (Color Tube) workbook -- Color Tube + kg-per-cone by Partita,
# used for the Biglietti "Color Tube" and "VMM22" columns. This is a
# different file from the raw-yarn Magazino used elsewhere in the app.
# ---------------------------------------------------------------------------

def load_vmm22_ratio_from_magazino(path: Path) -> tuple[dict[int, float], list[str]]:
    """{PARTITA: kg-per-cone} for VMM22, from the same raw Magazino Filato
    export already uploaded elsewhere in the app (Magazino Filato /
    Ordine Kamal tabs, cached path in magazino_cache.py) -- no separate
    upload needed. Per Partita: sum ESISTENZA (kg) and COLLI (cones)
    across MAGAZZINO 900910, 900160 and 900923, ratio = kg per cone.
    VMM22 for an order line = that ratio * the line's own Rocche."""
    errors: list[str] = []
    totals: dict[int, list[float]] = {}  # partita -> [esistenza_sum, colli_sum]
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    try:
        for row in _read_sheet_rows(wb.active):
            magazzino = _number(_get(row, "MAGAZZINO"))
            if magazzino not in (900910, 900160, 900923):
                continue
            partita = _number(_get(row, "PARTITA"))
            if partita is None:
                continue
            esistenza = _number(_get(row, "ESISTENZA")) or 0
            colli = _number(_get(row, "COLLI")) or 0
            bucket = totals.setdefault(int(partita), [0.0, 0.0])
            bucket[0] += esistenza
            bucket[1] += colli
    finally:
        wb.close()
    out = {p: (e / c) for p, (e, c) in totals.items() if c}
    if not out:
        errors.append("Nessuna riga valida trovata nel file Magazino Filato per VMM22.")
    return out, errors


@dataclass
class OrderRecord:
    customer_code: str
    customer_name: str
    article: str
    description: str
    additional_raw: str
    color_code: str
    color_name: str
    order_no: str
    order_row: str
    colored_batch: str
    raw_batch: str
    quantity_cones: float | int | None
    raw_weight: float | int | None
    dispo: str
    bagno: str
    delivery: Any = None
    machine: str = ""
    title: str = ""
    formato: str = "7777"
    commento: str = ""
    partita_med: str = ""
    cliente_med: str = ""
    kg: float | int | None = None
    densita: Any = ""
    color_tube: str = ""
    vmm22: float | int | None = None
    prezzo: Any = ""
    delivery_date: Any = ""


def _read_sheet_rows(ws) -> list[dict[str, Any]]:
    """Read a sheet using the best header row (row 1 or row 2)."""
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []
    candidates = []
    for idx in (0, 1):
        if idx < len(rows):
            headers = [_key(x) or f"__col{i + 1}" for i, x in enumerate(rows[idx])]
            score = sum(bool(x) for x in headers)
            candidates.append((score, idx, headers))
    # Prefer the first row when scores tie; a data row can contain as many
    # non-empty cells as a header row in ERP exports.
    _, header_idx, headers = max(candidates, key=lambda item: (item[0], -item[1]), default=(0, 0, []))
    out = []
    for values in rows[header_idx + 1:]:
        if not any(_clean(x) for x in values):
            continue
        # Many ERP exports repeat the header row twice.  Do not turn that
        # repeated header into a real data record (especially important for
        # the first Dispo-Bagno row).
        comparable = [(i, _key(values[i]), headers[i]) for i in range(min(len(values), len(headers))) if _clean(values[i]) and headers[i].startswith("__") is False]
        if comparable and sum(a == b for _, a, b in comparable) >= max(2, len(comparable) // 2):
            continue
        rec = {}
        for i, value in enumerate(values):
            if i < len(headers):
                # Keep the first duplicate header; this avoids row-2 blank
                # labels replacing the actual row-1 label.
                rec.setdefault(headers[i], value)
                rec.setdefault(f"__col{i + 1}", value)
        out.append(rec)
    return out


def _get(row: dict[str, Any], *names: str) -> Any:
    for name in names:
        if _key(name) in row and _clean(row[_key(name)]):
            return row[_key(name)]
    return ""


def detect_order_format(data_path: Path) -> str:
    """'ELVY_MED' or 'EL_KAMAL', by peeking at Sheet1's own header row --
    lets a single Data Ordine picker serve every customer instead of a
    separate one per source shape. EL KAMAL's Sheet1 already carries
    computed columns (CODICE/TITOLO/M-C/Clienti) that ELVY/MED's raw
    export never has; ELVY/MED's Sheet1 has 'Descrizione aggiuntiva
    ordine' and a numeric 'Cliente' code that EL KAMAL's never has.
    Defaults to 'ELVY_MED' if neither signature is conclusive."""
    wb = openpyxl.load_workbook(data_path, read_only=True)
    try:
        if "Sheet1" not in wb.sheetnames:
            return "ELVY_MED"
        rows = list(wb["Sheet1"].iter_rows(values_only=True, max_row=1))
    finally:
        wb.close()
    if not rows:
        return "ELVY_MED"
    headers = {_key(v) for v in rows[0] if v}
    el_kamal_signature = {_key("CODICE"), _key("Clienti"), _key("M/C")}
    elvy_med_signature = {_key("Descrizione aggiuntiva ordine"), _key("Cliente")}
    if el_kamal_signature.issubset(headers):
        return "EL_KAMAL"
    if elvy_med_signature.issubset(headers):
        return "ELVY_MED"
    return "EL_KAMAL" if "Doispo-Bagno" not in wb.sheetnames else "ELVY_MED"


def load_dispo_bagno_rows(path: Path) -> list[dict[str, Any]]:
    """Reads a Dispo-Bagno source regardless of shape: a .csv (EL KAMAL's
    own export), or an .xlsx with the data on its first sheet."""
    if path.suffix.lower() == ".csv":
        return _read_dispo_csv(path)
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    try:
        return _read_sheet_rows(wb.active)
    finally:
        wb.close()


def load_order(input_path: Path, dispo_path: Path | None = None) -> tuple[list[OrderRecord], list[dict[str, Any]]]:
    wb = openpyxl.load_workbook(input_path, data_only=True, read_only=True)
    if "Sheet1" not in wb.sheetnames:
        raise ValueError('Il file Data Ordine deve contenere il foglio "Sheet1".')
    data = _read_sheet_rows(wb["Sheet1"])
    dispo_rows = _read_sheet_rows(wb["Doispo-Bagno"]) if "Doispo-Bagno" in wb.sheetnames else []
    if dispo_path and dispo_path != input_path:
        dispo_rows = load_dispo_bagno_rows(dispo_path)
    raw_rows = _read_sheet_rows(wb["تحضير خيط خام"]) if "تحضير خيط خام" in wb.sheetnames else []
    if not data:
        raise ValueError("Sheet1 non contiene righe d'ordine.")
    dispo_by_riga = {
        _dispo_suffix(_get(r, "Dispo")): r
        for r in dispo_rows if _clean(_get(r, "Dispo"))
    }
    raw_by_article = {_clean(_get(r, "Articolo")).upper(): r for r in raw_rows}
    result: list[OrderRecord] = []
    for row in data:
        code = _clean(_get(row, "Cliente"))
        if code not in {"3009", "3004"}:
            continue
        article = _clean(_get(row, "Articolo"))
        raw = raw_by_article.get(("G" + article[1:]).upper(), {}) if article else {}
        additional = _clean(_get(row, "Descrizione aggiuntiva ordine"))
        raw_batch = _clean(_get(raw, "Partita.GG", "Partita")) or _first_match(additional, r"PG[- ]*([0-9]+)")
        riga = _clean(_get(row, "Riga"))
        dispo = dispo_by_riga.get(_dispo_suffix(riga), {})
        result.append(OrderRecord(
            customer_code=code,
            customer_name=_clean(_get(row, "__col2")) or ("MED" if code == "3004" else "ELVY"),
            article=article,
            description=additional or article,
            additional_raw=additional,
            color_code=_clean(_get(row, "Colore")),
            color_name=_clean(_get(row, "__col7")) or _clean(_get(row, "Colore")),
            order_no=_order_number(_get(row, "Ordine")),
            order_row=riga,
            colored_batch=_clean(_get(row, "Partita")),
            raw_batch=raw_batch,
            quantity_cones=_number(_get(row, "Ordinata", "Assegnata")) or _number(_get(raw, "عدد")),
            raw_weight=_number(_get(raw, "وزن", "Peso")),
            dispo=_clean(_get(dispo, "Dispo")),
            bagno=_clean(_get(dispo, "Field2", "Bagno")),
            delivery=_get(row, "Consegna"),
            title=_clean(_get(raw, "Titolo")) or _clean(_get(row, "Descrizione")),
        ))
    wb.close()
    if not result:
        raise ValueError("Non sono state trovate righe con Cliente 3004 o 3009.")
    return result, raw_rows


# ---------------------------------------------------------------------------
# EL KAMAL -- third customer, different source shape entirely: their own
# order export (Sheet1) already arrives with Titolo/KG/M-C/Bagno computed,
# instead of the raw ERP dump ELVY/MED use. The companion Dispo-Bagno file
# is a semicolon CSV (same columns as the Doispo-Bagno sheet) rather than a
# second sheet in the same workbook. Sheet1's own "Dispo" text embeds the
# row's sequence number as 'RIGA N', which is what joins to the CSV's
# 'D-...-00N' suffix -- same join pattern as MED's Riga column, just
# extracted from a different place.
# ---------------------------------------------------------------------------

def _read_dispo_csv(path: Path) -> list[dict[str, Any]]:
    import csv
    with open(path, encoding="utf-8-sig", newline="") as f:
        sample = f.read(4096)
        f.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=";,\t")
        except csv.Error:
            dialect = csv.excel
            dialect.delimiter = ";"
        reader = csv.reader(f, dialect)
        rows = list(reader)
    if not rows:
        return []
    headers = [_key(x) or f"__col{i + 1}" for i, x in enumerate(rows[0])]
    out = []
    for values in rows[1:]:
        if not any(_clean(x) for x in values):
            continue
        rec: dict[str, Any] = {}
        for i, value in enumerate(values):
            if i < len(headers):
                rec.setdefault(headers[i], value)
                rec.setdefault(f"__col{i + 1}", value)
        out.append(rec)
    return out


def load_el_kamal_order(data_path: Path, dispo_path: Path | None = None) -> tuple[list[OrderRecord], list[dict[str, Any]]]:
    wb = openpyxl.load_workbook(data_path, data_only=True, read_only=True)
    if "Sheet1" not in wb.sheetnames:
        raise ValueError('Il file Data Ordine EL KAMAL deve contenere il foglio "Sheet1".')
    data = _read_sheet_rows(wb["Sheet1"])
    wb.close()
    if not data:
        raise ValueError("Sheet1 non contiene righe d'ordine.")

    dispo_by_riga: dict[str, dict[str, Any]] = {}
    if dispo_path and dispo_path.is_file():
        dispo_rows = load_dispo_bagno_rows(dispo_path)
        dispo_by_riga = {
            _dispo_suffix(_get(r, "Dispo")): r
            for r in dispo_rows if _clean(_get(r, "Dispo"))
        }

    result: list[OrderRecord] = []
    for row in data:
        dispo_text = _clean(_get(row, "Dispo"))
        riga = _first_match(dispo_text, r"RIGA\s*([0-9]+)")
        dispo_row = dispo_by_riga.get(_dispo_suffix(riga), {}) if riga else {}
        bagno = _clean(_get(row, "Bagno")) or _clean(_get(dispo_row, "Field2", "Bagno"))
        result.append(OrderRecord(
            customer_code="3019",
            customer_name=_clean(_get(row, "Clienti")) or "EL KAMAL",
            article=_clean(_get(row, "CODICE", "Articolo")),
            description=_clean(_get(row, "COMMENTO RIGO ORDINE")),
            additional_raw=_clean(_get(row, "COMMENTO RIGO ORDINE")),
            color_code=_clean(_get(row, "COLORE", "Codice")),
            color_name=_clean(_get(row, "DESCR COL", "Colore")),
            order_no=_order_number(_get(row, "ORDINE", "Ordine")),
            order_row=riga,
            colored_batch=_clean(_get(row, "Partita Col", "Partita")),
            raw_batch=_clean(_get(row, "Partita GG")),
            quantity_cones=_number(_get(row, "ROCCHE", "Rocche")),
            raw_weight=_number(_get(row, "KG")),
            dispo=dispo_text or _clean(_get(dispo_row, "Dispo")),
            bagno=bagno,
            delivery=_get(row, "CONSEGNA", "Consegna"),
            title=_clean(_get(row, "TITOLO", "Titolo")),
            formato=_clean(_get(row, "Formmato", "Formato")) or "7777",
            machine=_clean(_get(row, "M/C")),
            commento=_clean(_get(row, "COMMENTO RIGO ORDINE")),
            kg=_number(_get(row, "KG")),
        ))
    if not result:
        raise ValueError("Non sono state trovate righe d'ordine EL KAMAL in Sheet1.")
    return result, []


def build_el_kamal_stem(records: list["OrderRecord"]) -> str:
    """e.g. 'D-505444-EL_KAMAL' -- confirmed against a real reference
    ticket filename the user sent (no year in it, unlike MED)."""
    dispo_num = _dispo_number(records[0].dispo) if records else ""
    return f"D-{dispo_num}-EL_KAMAL" if dispo_num else "EL_KAMAL"


def enrich_records(
    records: list[OrderRecord],
    customer: str,
    codes_map: dict[str, str] | None = None,
    densita_map: dict[int, dict[str, Any]] | None = None,
    vmm_ratio_map: dict[int, float] | None = None,
    price_lookup: dict[tuple, tuple] | None = None,
) -> None:
    """Fills in Titolo, Commento, M/C, KG, Densita, Color Tube, VMM22,
    Prezzo and the MED "Partita MED"/"Cliente MED" pair, in place.
    Anything whose source file was not provided is left blank, same as
    before, instead of raising.
    """
    codes_map = codes_map or {}
    densita_map = densita_map or {}
    vmm_ratio_map = vmm_ratio_map or {}
    price_lookup = price_lookup or {}

    # --- M/C: group by Bagno, sum Rocche (+ MED Polmoni bonus). EL KAMAL's
    # own source already gives a correct M/C per row, so it's left alone.
    if customer != "EL_KAMAL":
        groups: dict[str, list[OrderRecord]] = {}
        for r in records:
            groups.setdefault(r.bagno or "", []).append(r)
        for bagno, members in groups.items():
            total = sum(_number(m.quantity_cones) or 0 for m in members)
            if customer == "MED":
                total += sum(_polmoni_extra(m.additional_raw) for m in members)
            extra = MACHINE_SIZE_TABLE_MED_EXTRA if customer == "MED" else None
            machine = _machine_for_count(total, extra=extra)
            for m in members:
                m.machine = machine

    for r in records:
        # Titolo: Articoli.xlsx lookup takes priority (fuller description);
        # the raw yarn sheet's shorter Titolo is the fallback.
        r.title = _titolo_lookup(r.article, codes_map) or r.title

        # EL KAMAL's Commento (COMMENTO RIGO ORDINE) already arrives correct
        # from the source and isn't a dash-encoded comment like ELVY/MED's.
        if customer != "EL_KAMAL":
            r.commento = _commento_for(customer, r.additional_raw)

        if customer == "MED":
            r.partita_med = _partita_med(r.additional_raw)
            r.cliente_med = _cliente_med(r.additional_raw)

        try:
            batch_key = int(_number(r.raw_batch)) if _number(r.raw_batch) is not None else None
        except (TypeError, ValueError):
            batch_key = None

        if batch_key is not None and batch_key in densita_map:
            entry = densita_map[batch_key]
            peso_net = entry.get("peso_net")
            r.kg = int(round(peso_net * (_number(r.quantity_cones) or 0))) if peso_net is not None else r.raw_weight
            r.densita = entry.get("densita", "")
            r.color_tube = entry.get("color_tube", "")
        else:
            r.kg = r.raw_weight

        if batch_key is not None and batch_key in vmm_ratio_map:
            kg_per_cone = vmm_ratio_map[batch_key]
            r.vmm22 = round(kg_per_cone * (_number(r.quantity_cones) or 0), 2)

        r.prezzo = _prezzo_for(r.article, r.color_code, price_lookup)

    if customer == "ELVY":
        compute_delivery_date(records)


def _filato_rows(records: list["OrderRecord"], raw_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """One row per raw batch (Partita GG), Rocche summed from the actual
    order records -- not read off the small 'تحضير خيط خام' reference
    sheet's own count, which can be stale/incomplete on its own (a real
    case: a raw batch's order Rocche was 72 but that sheet's count showed
    0 for it). Titolo/Peso still come from that sheet when available,
    purely for context."""
    raw_by_article = {_clean(_get(r, "Articolo")).upper(): r for r in raw_rows}

    totals: dict[str, float] = {}
    for rec in records:
        key = _clean(rec.raw_batch)
        if not key:
            continue
        totals[key] = totals.get(key, 0) + (_number(rec.quantity_cones) or 0)

    out = []
    seen: set[str] = set()
    for rec in records:
        key = _clean(rec.raw_batch)
        if not key or key in seen:
            continue
        seen.add(key)
        article_g = ("G" + rec.article[1:]) if rec.article[:1].upper() == "C" else rec.article
        raw = raw_by_article.get(article_g.upper(), {})
        out.append({
            "Articolo": article_g,
            "Titolo": rec.title or _clean(_get(raw, "Titolo")),
            "Partita": key,
            "Rocche": totals[key],
            "Peso": _number(_get(raw, "وزن", "Peso")),
            "تحضير خام": _clean(_get(raw, "Custom", "تحضير خام")) or "تحضير خام",
        })
    return out


def export_workbook(
    path: Path,
    records: list[OrderRecord],
    raw_rows: list[dict[str, Any]],
    include_filato: bool = True,
    stem: str = "",
    customer: str = "",
) -> None:
    customer = customer or ("ELVY" if records[0].customer_code == "3009" else "MED")
    wb = Workbook()
    ws = wb.active
    ws.title = customer
    common = ["Dispo/Riga", "Articolo", "Titolo", "Formato", "Ordine", "Codice", "Colore", "Rocche", "KG", "M/C", "Partita Col", "Consegna", "Commento", "Bagno", "Cliente", "Partita GG"]
    extra_cols = ["Color Tube", "VMM22", "Prezzo", "Densita` (360-390)"]
    if customer == "MED":
        headers = common + ["Partita MED", "Cliente MED", "POLMON"] + extra_cols
    elif customer == "EL_KAMAL":
        # EL KAMAL's source doesn't carry these -- dropped per request, not
        # just left blank.
        headers = common
    else:
        headers = common + ["Delivery Date"] + extra_cols
    ws.append(headers)
    for r in records:
        row = [r.dispo, r.article, r.title, r.formato, r.order_no, r.color_code, r.color_name,
               r.quantity_cones, r.kg, r.machine, r.colored_batch, r.delivery,
               r.commento, r.bagno, r.customer_name, r.raw_batch]
        if customer == "MED":
            polmon = _polmoni_segment(r.additional_raw)
            row += [r.partita_med, r.cliente_med, polmon, r.color_tube, r.vmm22, r.prezzo, r.densita]
        elif customer == "EL_KAMAL":
            pass
        else:
            row += [r.delivery_date, r.color_tube, r.vmm22, r.prezzo, r.densita]
        ws.append(row)
    if include_filato:
        fws = wb.create_sheet("Filato x Tinturia")
        fheaders = ["Articolo", "Titolo", "Partita", "Rocche", "Peso", "تحضير خام"]
        fws.append(fheaders)
        for r in _filato_rows(records, raw_rows):
            fws.append([r[h] for h in fheaders])
        _style_sheet(fws)
    _style_sheet(
        ws,
        date_columns=("Consegna", "Delivery Date"),
        duplicate_highlight_columns=("Bagno",),
        range_highlight_columns={"Densita` (360-390)": (360, 390)},
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)
    wb.close()


def export_filato_workbook(path: Path, records: list["OrderRecord"], raw_rows: list[dict[str, Any]]) -> None:
    """Export only the optional ``Filato x Tinturia`` workbook."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Filato x Tinturia"
    headers = ["Articolo", "Titolo", "Partita", "Rocche", "Peso", "تحضير خام"]
    ws.append(headers)
    for r in _filato_rows(records, raw_rows):
        ws.append([r[h] for h in headers])
    _style_sheet(ws)
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)
    wb.close()


def _style_sheet(
    ws,
    date_columns: tuple[str, ...] = (),
    duplicate_highlight_columns: tuple[str, ...] = (),
    range_highlight_columns: dict[str, tuple[float, float]] | None = None,
) -> None:
    """Bold/blue header, every cell (not just the header) centered, optional
    date-only number format for the given columns, optional red highlight
    for duplicate values in the given columns (e.g. Bagno), and optional
    red highlight for numbers outside a (min, max) range in the given
    columns (e.g. Densita`(360-390))."""
    fill = PatternFill("solid", fgColor="FF1F4E78")
    header = [c.value for c in ws[1]]
    last_row = ws.max_row

    thin = Side(style="thin", color="000000")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    band_fill = PatternFill("solid", fgColor="FFDCE6F1")
    for row in ws.iter_rows():
        banded = row[0].row > 1 and (row[0].row % 2 == 0)
        for cell in row:
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=(cell.row == 1))
            cell.border = border
            if banded:
                cell.fill = band_fill
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = fill

    for col_name in date_columns:
        if col_name in header:
            col_idx = header.index(col_name) + 1
            for r in range(2, last_row + 1):
                ws.cell(row=r, column=col_idx).number_format = "dd/mm/yyyy"

    if last_row > 1:
        from openpyxl.formatting.rule import FormulaRule
        # Differential-format (conditional-formatting) fills render from
        # bgColor, not fgColor -- a plain PatternFill("solid", fgColor=...)
        # writes only fgColor into the dxf and Excel shows no highlight at
        # all even though the rule fires. Setting both covers it either way.
        red_fill = PatternFill(start_color="FFFFC7CE", end_color="FFFFC7CE", fill_type="solid")

        for col_name in duplicate_highlight_columns:
            if col_name in header:
                col_letter = ws.cell(row=1, column=header.index(col_name) + 1).column_letter
                rng = f"{col_letter}2:{col_letter}{last_row}"
                ws.conditional_formatting.add(
                    rng,
                    FormulaRule(
                        formula=[f'AND(${col_letter}2<>"",COUNTIF(${col_letter}$2:${col_letter}${last_row},{col_letter}2)>1)'],
                        fill=red_fill,
                        stopIfTrue=False,
                    ),
                )

        for col_name, (lo, hi) in (range_highlight_columns or {}).items():
            if col_name in header:
                col_letter = ws.cell(row=1, column=header.index(col_name) + 1).column_letter
                rng = f"{col_letter}2:{col_letter}{last_row}"
                first = f"{col_letter}2"
                formula = f'AND({first}<>"",OR({first}<{lo},{first}>{hi}))'
                ws.conditional_formatting.add(rng, FormulaRule(formula=[formula], fill=red_fill, stopIfTrue=False))

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    for col in ws.columns:
        letter = col[0].column_letter
        content_width = max((len(_clean(c.value)) for c in col), default=8)
        ws.column_dimensions[letter].width = max(10, content_width + 2)


def export_word(path: Path, template_path: Path, records: list[OrderRecord], stem: str = "") -> None:
    """Create one Word file, one Biglietto page per order color.

    Uses python-docx (lxml-backed) instead of the stdlib ElementTree, and
    deep-copies the template's own table element rather than re-parsing a
    serialized copy. Both changes avoid the namespace-prefix rewriting that
    xml.etree.ElementTree does on round-trip, which is the most likely
    cause of Word's "unreadable content" repair prompt on the old output.
    """
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches

    doc = Document(str(template_path))
    if not doc.tables:
        raise ValueError("Il modello Biglietti non contiene la tabella prevista.")

    # The template's own page margins (1in top/bottom) leave the ticket
    # table only ~0.05in of headroom before it spills its last 2-3 rows
    # onto a second page whenever a bilingual label wraps to 2 lines --
    # that's the "biglietto finishes outside its own page" issue. Trimming
    # the margins buys back room without touching row heights or fonts.
    for sect in doc.sections:
        sect.top_margin = Inches(0.5)
        sect.bottom_margin = Inches(0.5)

    template_tbl = doc.tables[0]._tbl
    body = doc.element.body
    template_tbl.getparent().remove(template_tbl)
    sect_pr = body.find(qn("w:sectPr"))

    def _insert(el) -> None:
        if sect_pr is not None:
            sect_pr.addprevious(el)
        else:
            body.append(el)

    for idx, record in enumerate(records):
        new_tbl = copy.deepcopy(template_tbl)
        _fill_ticket_table(new_tbl, record, qn, OxmlElement, stem)
        _insert(new_tbl)
        if idx != len(records) - 1:
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            r.append(br)
            p.append(r)
            _insert(p)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def _fill_ticket_table(tbl, r: OrderRecord, qn, OxmlElement, stem: str) -> None:
    rows = tbl.findall(qn("w:tr"))

    def put(row: int, cell: int, text: Any) -> None:
        c = rows[row - 1].findall(qn("w:tc"))[cell - 1]
        for p in c.findall(".//" + qn("w:p")):
            for t in p.findall(".//" + qn("w:t")):
                t.text = ""
            first = p.find(".//" + qn("w:t"))
            if first is None:
                run = OxmlElement("w:r")
                first = OxmlElement("w:t")
                run.append(first)
                p.append(run)
            first.text = _clean(text)
            break

    put(1, 2, _ticket_header(r, stem))
    put(2, 2, r.title); put(2, 3, r.formato); put(2, 4, r.article)
    put(3, 2, r.raw_batch); put(4, 2, r.color_name); put(5, 2, r.colored_batch)
    put(6, 2, r.dispo); put(6, 3, f"Ordine.  {r.order_no}"); put(7, 3, r.bagno)
    put(8, 2, r.quantity_cones); put(9, 2, r.kg); put(10, 2, r.machine)
    vmm = r.vmm22 if r.vmm22 is not None else r.raw_weight
    put(11, 3, f"VMM22( {float(vmm):.2f} )Kg" if vmm is not None else "VMM22(       )Kg")


def _format_date(value: Any) -> str:
    try:
        return value.strftime("%d/%m/%Y")
    except AttributeError:
        return _clean(value)


def _ticket_header(r: OrderRecord, stem: str) -> str:
    """Top line of the ticket. Differs per customer:
    - MED: first 3 letters of Cliente - Commento - Polmon - Cliente MED
      (blank segments dropped rather than left as bare dashes).
    - EL KAMAL: Cliente - Consegna.
    - Everything else (ELVY): Cliente - <order identifier>, unchanged."""
    if r.customer_code == "3004":  # MED
        polmon = _polmoni_segment(r.additional_raw)
        parts = [r.customer_name[:3], r.commento, polmon, r.cliente_med]
        return "-".join(p for p in parts if p)
    if r.customer_code == "3019":  # EL KAMAL
        return f"{r.customer_name} - {_format_date(r.delivery)}"
    return f"{r.customer_name}-{stem}"
